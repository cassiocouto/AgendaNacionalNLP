from docx import Document
import pandas as pd

stopwords = [
    "da",
    "de",
    "e",
    "em",
    "no",
    "na",
    "com",
    "para",
    "os",
    "as",
    "um",
    "uma",
    "ou",
    "que",
    "se",
    "são",
    "por",
    "dos",
    "das",
    "nos",
    "nas",
    "como",
    "mais",
    "sobre",
    "o",
    "a",
    "os",
    "as",
]

regions = ["NORTE", "NORDESTE", "CENTRO-OESTE", "SUDESTE", "SUL"]
states = [
    "ACRE",
    "ALAGOAS",
    "AMAPÁ",
    "AMAZONAS",
    "BAHIA",
    "CEARÁ",
    "DISTRITO FEDERAL",
    "ESPÍRITO SANTO",
    "GOIÁS",
    "MARANHÃO",
    "MATO GROSSO",
    "MATO GROSSO DO SUL",
    "MINAS GERAIS",
    "PARÁ",
    "PARAÍBA",
    "PARANÁ",
    "PERNAMBUCO",
    "PIAUÍ",
    "RIO DE JANEIRO",
    "RIO GRANDE DO NORTE",
    "RIO GRANDE DO SUL",
    "RONDÔNIA",
    "RORAIMA",
    "SANTA CATARINA",
    "SÃO PAULO",
    "SERGIPE",
    "TOCANTINS",
]

ignored_phrases = ["Declaração de Prioridades:"]

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)


def clean_stopwords(text, stopwords):
    text = " ".join([word for word in text.split() if word.lower() not in stopwords])
    return text

def clean_ignore_phrases(text, ignored_phrases):
    for ignored_phrase in ignored_phrases:
        text = text.replace(ignored_phrase, "")
    return text


def read_themes_from_docx(file_path):
    doc = Document(file_path)
    curr_region = ""
    curr_state = ""
    region_arr = []
    state_arr = []
    theme_arr = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in regions:
            curr_region = text
        elif text in states:
            curr_state = text
        elif text and text not in ignored_phrases:
            region_arr.append(curr_region)
            state_arr.append(curr_state)
            theme_arr.append(text)

    data = pd.DataFrame({"region": region_arr, "state": state_arr, "theme": theme_arr})

    return data

def preprocess_text_and_tokenize(df, stopwords):
    df["cleaned_theme"] = df["theme"].apply(lambda x: clean_stopwords(x, stopwords))
    return df